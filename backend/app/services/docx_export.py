from typing import Dict, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.services.template_parser import parse_template_structure


def create_paper_from_sections(
    sections: Dict[str, str],
    output_path: str,
    template_path: Optional[str] = None
) -> str:
    """
    Create a DOCX paper from sections dictionary.
    
    Args:
        sections: Dictionary mapping section titles to content
        output_path: Path where the output DOCX will be saved
        template_path: Optional template file to match structure
        
    Returns:
        Path to the created DOCX file
    """
    # Initialize document
    if template_path and Path(template_path).exists():
        # Use template structure
        template_structure = parse_template_structure(template_path)
        doc = Document()
        _apply_sections_with_template(doc, sections, template_structure)
    else:
        # Create document without template
        doc = Document()
        _apply_sections_without_template(doc, sections)
    
    # Save document
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    
    return str(output_file)


def _apply_sections_with_template(
    doc: Document,
    sections: Dict[str, str],
    template_structure: Dict
) -> None:
    """
    Apply sections to document following template structure.
    
    Args:
        doc: Document object to populate
        sections: Section content dictionary
        template_structure: Template structure from parser
    """
    section_order = template_structure.get("section_order", [])
    headings_info = {h["text"]: h for h in template_structure.get("headings", [])}
    
    for section_title in section_order:
        content = sections.get(section_title, "")
        heading_info = headings_info.get(section_title, {})
        level = heading_info.get("level", 1)
        
        # Add heading
        heading_style = f"Heading {level}"
        doc.add_heading(section_title, level=level)
        
        # Add content
        if content:
            _add_formatted_content(doc, content)


def _apply_sections_without_template(
    doc: Document,
    sections: Dict[str, str]
) -> None:
    """
    Apply sections to document without template (simple structure).
    
    Args:
        doc: Document object to populate
        sections: Section content dictionary
    """
    for section_title, content in sections.items():
        # Add heading
        doc.add_heading(section_title, level=1)
        
        # Add content
        if content:
            _add_formatted_content(doc, content)


def _add_formatted_content(doc: Document, content: str) -> None:
    """
    Add formatted content to document, handling paragraphs and line breaks.
    
    Args:
        doc: Document object
        content: Text content to add
    """
    # Split content into paragraphs (double line breaks)
    paragraphs = content.split("\n\n")
    
    for para_text in paragraphs:
        if para_text.strip():
            # Check if it's a bullet point or numbered list
            if para_text.strip().startswith("- ") or para_text.strip().startswith("* "):
                # Handle bullet lists
                items = [line.strip()[2:] for line in para_text.split("\n") 
                        if line.strip().startswith(("- ", "* "))]
                for item in items:
                    doc.add_paragraph(item, style='List Bullet')
            elif _is_numbered_list(para_text):
                # Handle numbered lists
                items = [line.strip().split(". ", 1)[1] for line in para_text.split("\n") 
                        if line.strip() and ". " in line]
                for item in items:
                    doc.add_paragraph(item, style='List Number')
            else:
                # Regular paragraph
                paragraph = doc.add_paragraph(para_text.strip())
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _is_numbered_list(text: str) -> bool:
    """Check if text starts with a numbered list pattern."""
    first_line = text.strip().split("\n")[0].strip()
    return first_line and first_line[0].isdigit() and ". " in first_line


def export_with_citations(
    sections: Dict[str, str],
    citations: list,
    output_path: str,
    template_path: Optional[str] = None
) -> str:
    """
    Create a DOCX paper with sections and citations.
    
    Args:
        sections: Dictionary mapping section titles to content
        citations: List of citation dictionaries
        output_path: Path where the output DOCX will be saved
        template_path: Optional template file to match structure
        
    Returns:
        Path to the created DOCX file
    """
    # Create base document with sections
    output_file = create_paper_from_sections(sections, output_path, template_path)
    
    # Add citations if provided
    if citations:
        doc = Document(output_file)
        
        # Add references section
        doc.add_page_break()
        doc.add_heading("References", level=1)
        
        for citation in citations:
            if isinstance(citation, dict):
                citation_text = citation.get("text", str(citation))
            else:
                citation_text = str(citation)
            
            paragraph = doc.add_paragraph(citation_text)
            # Format as hanging indent
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        
        # Save updated document
        doc.save(output_file)
    
    return output_file


def merge_sections(
    base_sections: Dict[str, str],
    additional_sections: Dict[str, str]
) -> Dict[str, str]:
    """
    Merge two sections dictionaries, with additional_sections taking precedence.
    
    Args:
        base_sections: Base sections dictionary
        additional_sections: Additional sections to merge/override
        
    Returns:
        Merged sections dictionary
    """
    merged = base_sections.copy()
    merged.update(additional_sections)
    return merged
